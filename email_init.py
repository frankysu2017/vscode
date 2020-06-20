#!/usr/bin/env python3
# coding=utf-8
# email_init.py

import email
import re
import os
import shutil
import chardet
from datetime import datetime
from bs4 import BeautifulSoup
from dateutil.parser import parse
from docx import Document
import win32com
import win32com.client
import xlrd, openpyxl
wc = win32com.client.constants


def get_subject(msg):
    s = msg.get('subject')
    if s:
        s_decoded = email.header.decode_header(s)
        if isinstance(s_decoded[0][0], bytes):
            return s_decoded[0][0].decode(s_decoded[0][1], 'ignore')
        else:
            return s_decoded[0][0]
    else:
        return None


def get_sendtime(msg):
    if 'date' in msg:
        return parse(msg.get('date'))
    else:
        return None


def get_sender(msg):
    sname, send_box = email.utils.parseaddr(msg.get('from'))
    if sname:
        sname = email.header.decode_header(sname)
        if isinstance(sname[0][0], bytes):
            send_name = sname[0][0].decode(sname[0][1], 'ignore')
        else:
            send_name = sname[0][0]
    else:
        send_name = None
    return send_name, send_box


def get_receiver(msg):
    rname, receive_box = email.utils.parseaddr(msg.get('to'))
    if rname:
        rname = email.header.decode_header(rname)
        if isinstance(rname[0][0], bytes):
            receive_name = rname[0][0].decode(rname[0][1], 'ignore')
        else:
            receive_name = rname[0][0]
    else:
        receive_name = rname
    return receive_name, receive_box


def get_mainbody(msg):
    mainbody = ''
    if msg.get_payload(decode=True) and msg.get_content_charset():
        mainbody = msg.get_payload(decode=True).strip().decode(msg.get_content_charset(), 'ignore')
        html_flag = re.search('text/(.*)', msg.get_content_type()).group(1).lower()
        if html_flag == 'html':
            if '<html' in mainbody:
                prefix, html_text, suffix = re.search('(.*)(<html.*</html>)(.*)', mainbody, re.DOTALL+re.IGNORECASE).groups()
                html_text = BeautifulSoup(html_text, features="html.parser").body.get_text()
                mainbody = prefix + '\n' + html_text + '\n' + suffix
    mainbody = re.sub('[\n]+', '\n', mainbody)
    return mainbody


def read_excel(filename):
    excel = xlrd.open_workbook(filename)
    attachment_text = ''
    for sheet in excel.sheets():
        for row in sheet.get_rows():
            attachment_text += ','.join([str(cell.value) for cell in row])
        attachment_text += '\n\n\n'
    return attachment_text


attachment_file_path = 'C:\\Users\\WEB_Station\\PycharmProjects\\hkprofile\\attachments\\'
def read_word(filename):
    extension_name = filename.split('.')[-1].lower()
    doc = None
    if extension_name == 'doc':
        try:
            wps = win32com.client.gencache.EnsureDispatch('word.application')
        except:
            wps = win32com.client.gencache.EnsureDispatch('wps.application')
        else:
            wps = win32com.client.gencache.EnsureDispatch('word.application')
        wps.Visible = False
        d = wps.Documents.Open(filename)
        try:
            d.SaveAs2(attachment_file_path + 'temp.docx', 12)
        except:
            pass
        try:
            wps.Documents.Close()
            wps.Documents.Close(wc.wdDoNotSaveChanges)
            wps.Quit
        except:
            pass
        try:
            doc = Document(attachment_file_path + r'temp.docx')
        except:
            doc = Document(attachment_file_path + r'temp2.docx')
    else:
        try:
            doc = Document(filename)
        except:
            pass
    attachemnt_text = ''
    if doc:
        for each in doc.paragraphs:
            attachemnt_text += each.text + '\n'
    return attachemnt_text


def get_attachment(msg, mail_address):
    attachemnt_name = msg.get_filename()
    if attachemnt_name:
        attachemnt_name = email.header.decode_header(attachemnt_name)
        if isinstance(attachemnt_name[0][0], bytes):
            attachemnt_name = attachemnt_name[0][0].decode(attachemnt_name[0][1], 'ignore')
        else:
            attachemnt_name = attachemnt_name[0][0]
        if '\\' in attachemnt_name:
            attachemnt_name = attachemnt_name.replace('\\', '_').replace(':', '')

        attachemnt_content = msg.get_payload(decode=True).strip()
        path = attachment_file_path + mail_address + '\\'
        if not os.path.exists(path):
            os.mkdir(path)
        with open(path + attachemnt_name, 'wb') as f:
            f.write(attachemnt_content)
        extension_name = attachemnt_name.split('.')[-1].lower()
        if extension_name in ['doc', 'docx']:
            attachment_text = read_word(path + attachemnt_name)
            return attachemnt_name, attachment_text
        elif extension_name in ['xls', 'xlsx']:
            attachment_text = read_excel(path + attachemnt_name)
            return attachemnt_name, attachment_text
        else:
            return attachemnt_name, ''


def email_init(emlfile):
    #print(emlfile)
    mail_address = emlfile.split('/')[-1].split('_')[0]
    encode = chardet.detect(open(emlfile, 'rb').read())['encoding']
    with open(emlfile, 'r', encoding=encode) as eml:
        msg = email.message_from_file(eml)
        subject = get_subject(msg)
        send_time = get_sendtime(msg)
        send_name, send_box = get_sender(msg)
        receive_name, receive_box = get_receiver(msg)
        attachment_file, attachment_text = '', ''
        attachment_file_temp, attachment_text_temp = '', ''
        if msg.is_multipart():
            #print('it is multipart:\n')
            mail_content = ''
            for part in msg.get_payload():
                if part.get_content_maintype() == 'multipart':
                    part_content = ''
                    for subpart in part.get_payload():
                        subpart_content = get_mainbody(subpart)
                        part_content += subpart_content
                elif part.get_content_maintype() == 'text':
                    part_content = get_mainbody(part)
                elif part.get_content_maintype() == 'application':
                    part_content = ''
                    attachment_file_temp, attachment_text_temp = get_attachment(part, mail_address)
                else:
                    part_content = ''
                    attachment_file_temp, attachment_text_temp = get_attachment(part, mail_address)
                attachment_file = attachment_file_temp + '\n' + attachment_file
                attachment_text = attachment_text_temp + '\n' + attachment_text
                mail_content += part_content
        else:
            #print('it is text: \n')
            mail_content = get_mainbody(msg)

        return send_name, send_box, receive_name, receive_box, str(send_time), subject, mail_content, attachment_file, attachment_text


if __name__ == '__main__':
    file = openpyxl.Workbook()
    table = file.create_sheet('data')
    q = 0
    start = datetime.now()
    ILLEGAL_CHARACTERS_RE = re.compile(r'[\000-\010]|[\013-\014]|[\016-\037]')
    for i, item in enumerate(os.listdir(r'D:\email/')):
        try:
            r = email_init(r'D:\email/{}'.format(item))
            for j, value in enumerate(r):
                if isinstance(value, str):
                    value = ILLEGAL_CHARACTERS_RE.sub(r'', value)
                table.cell(row=i+1, column=j+1).value = value
        except:
            shutil.copyfile(r'D:\email/{}'.format(item), r'./error/{}'.format(item))
    print('processing costs {} seconds'.format(datetime.now() - start))
    file.save(r'./data1.xlsx')

